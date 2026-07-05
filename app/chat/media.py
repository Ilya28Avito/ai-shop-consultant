import base64
import tempfile
import os
import subprocess
from io import BytesIO
from fastapi import UploadFile
from openai import AsyncOpenAI
from dotenv import load_dotenv

load_dotenv(".env_robust_23")
client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))


async def media_to_part(media: UploadFile) -> dict:
    """Конвертирует медиафайл в content-part для OpenAI."""
    mime = media.content_type or ""
    data = await media.read()

    if mime.startswith("image/"):
        b64 = base64.b64encode(data).decode()
        return {
            "type": "image_url",
            "image_url": {"url": f"data:{mime};base64,{b64}"}
        }

    if mime.startswith("audio/") or mime == "application/ogg":
        transcript = await whisper_transcribe(data)
        return {
            "type": "text",
            "text": f"[пользователь сказал голосом]:\n{transcript}"
        }

    if mime == "application/pdf":
        return {
            "type": "text",
            "text": f"[документ PDF]:\n{extract_pdf_text(data)[:30_000]}"
        }

    if mime.endswith("wordprocessingml.document"):
        return {
            "type": "text",
            "text": f"[документ DOCX]:\n{extract_docx_text(data)[:30_000]}"
        }

    raise ValueError(f"Unsupported media type: {mime}")


async def whisper_transcribe(audio_bytes: bytes) -> str:
    """Транскрибирует аудио через Whisper-1 с конвертацией через ffmpeg."""
    with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp_ogg:
        tmp_ogg.write(audio_bytes)
        tmp_ogg_path = tmp_ogg.name

    tmp_mp3_path = tmp_ogg_path.replace(".ogg", ".mp3")

    try:
        subprocess.run(
            ["ffmpeg", "-y", "-i", tmp_ogg_path, tmp_mp3_path],
            capture_output=True, check=True
        )

        with open(tmp_mp3_path, "rb") as f:
            mp3_data = f.read()

        result = await client.audio.transcriptions.create(
            model="whisper-1",
            file=("voice.mp3", mp3_data, "audio/mpeg"),
        )
        return result.text
    finally:
        if os.path.exists(tmp_ogg_path):
            os.unlink(tmp_ogg_path)
        if os.path.exists(tmp_mp3_path):
            os.unlink(tmp_mp3_path)


def extract_pdf_text(data: bytes) -> str:
    """Извлекает текст из PDF."""
    try:
        import pypdf
        reader = pypdf.PdfReader(BytesIO(data))
        pages = reader.pages[:50]
        text = "\n".join(page.extract_text() or "" for page in pages)
        return text
    except Exception as e:
        return f"[Ошибка извлечения текста из PDF: {e}]"


def extract_docx_text(data: bytes) -> str:
    """Извлекает текст из DOCX."""
    try:
        import docx
        doc = docx.Document(BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                tables_text.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(paragraphs + tables_text)
    except Exception as e:
        return f"[Ошибка извлечения текста из DOCX: {e}]"
